"""Metadata extractor for enriching file summaries with detailed metadata.

This module provides programmatic and LLM-based extraction of file metadata
to support diverse task synthesis.
"""

import ast
import json
import os
from pathlib import Path
from typing import Any

from ..utils.llm_client import LLMClient


# File extension to category mapping
EXTENSION_CATEGORY_MAP = {
    # Table/Spreadsheet
    ".csv": "table",
    ".xlsx": "table",
    ".xls": "table",
    ".xlsm": "table",
    ".xltx": "table",
    ".tsv": "table",
    # Document
    ".pdf": "document",
    ".docx": "document",
    ".doc": "document",
    ".txt": "document",
    ".md": "document",
    ".rtf": "document",
    # Code
    ".py": "code",
    ".js": "code",
    ".ts": "code",
    ".java": "code",
    ".cpp": "code",
    ".c": "code",
    ".go": "code",
    ".rs": "code",
    ".rb": "code",
    ".php": "code",
    ".sh": "code",
    ".ipynb": "notebook",
    # Image
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".gif": "image",
    ".tiff": "image",
    ".bmp": "image",
    ".webp": "image",
    # Scientific
    ".fasta": "scientific",
    ".fa": "scientific",
    ".fastq": "scientific",
    ".fq": "scientific",
    ".vcf": "scientific",
    ".sam": "scientific",
    ".bam": "scientific",
    ".fits": "scientific",
    ".geojson": "scientific",
    # Presentation
    ".pptx": "presentation",
    ".ppt": "presentation",
    # Data
    ".json": "data",
    ".xml": "data",
    ".yaml": "data",
    ".yml": "data",
    ".html": "markup",
}


# LLM prompt for semantic metadata extraction
LLM_METADATA_EXTRACTION_PROMPT = """Based on the file information below, extract semantic metadata.

## File Information
- **Filename**: {filename}
- **Extension**: {extension}
- **File Size**: {file_size} bytes
- **Content Type**: {content_type}

## File Summary
{summary}

## Programmatic Metadata (already extracted)
{programmatic_metadata}

---

## Task

Extract the following semantic metadata as a JSON object:

1. **language**: Primary language of the document (e.g., "en", "zh", "mixed", "code")
2. **domain**: Domain classification (one of: "finance", "medical", "legal", "scientific", "government", "education", "business", "technology", "general")
3. **document_type**: Type of document (e.g., "report", "form", "invoice", "resume", "paper", "template", "statement", "contract", "manual", "data")
4. **complexity_level**: Complexity (one of: "simple", "moderate", "complex")
5. **key_entities**: List of up to 5 important named entities (organizations, products, concepts)
6. **potential_operations**: List of 3-5 operations that could be performed on this file (e.g., "extract_tables", "calculate_totals", "fill_form", "convert_format", "analyze_trends")

Return ONLY a valid JSON object, no explanation:
```json
{{
  "language": "...",
  "domain": "...",
  "document_type": "...",
  "complexity_level": "...",
  "key_entities": ["...", "..."],
  "potential_operations": ["...", "..."]
}}
```
"""


class MetadataExtractor:
    """Extracts detailed metadata from files using programmatic and LLM-based methods."""

    def __init__(self, llm_client: LLMClient | None = None):
        """
        Initialize the metadata extractor.

        Args:
            llm_client: Optional LLM client for semantic metadata extraction
        """
        self.llm = llm_client

    def extract(
        self,
        file_path: str,
        summary: str,
        content_type: str,
    ) -> dict[str, Any]:
        """
        Extract detailed metadata from a file.

        Args:
            file_path: Path to the file
            summary: Already generated file summary
            content_type: Already classified content type (form, text, table, code, mixed)

        Returns:
            Dictionary containing extracted metadata
        """
        metadata: dict[str, Any] = {}
        path = Path(file_path)
        ext = path.suffix.lower()

        # 1. Extract common metadata (programmatic)
        metadata.update(self._extract_common(file_path))

        # 2. Extract type-specific metadata (programmatic)
        category = EXTENSION_CATEGORY_MAP.get(ext, "unknown")
        type_metadata = self._extract_by_category(file_path, ext, category)
        metadata.update(type_metadata)

        # 3. Extract semantic metadata (LLM-based)
        if self.llm:
            try:
                semantic = self._extract_semantic(
                    file_path=file_path,
                    summary=summary,
                    content_type=content_type,
                    programmatic_metadata=metadata,
                )
                metadata.update(semantic)
            except Exception as e:
                metadata["llm_extraction_error"] = str(e)[:100]

        return metadata

    def _extract_common(self, file_path: str) -> dict[str, Any]:
        """Extract common metadata available for all files."""
        path = Path(file_path)
        metadata = {}

        try:
            stat = path.stat()
            metadata["file_size_bytes"] = stat.st_size
        except OSError:
            metadata["file_size_bytes"] = None

        return metadata

    def _extract_by_category(
        self,
        file_path: str,
        ext: str,
        category: str,
    ) -> dict[str, Any]:
        """Extract category-specific metadata."""
        extractors = {
            "table": self._extract_table_metadata,
            "document": self._extract_document_metadata,
            "code": self._extract_code_metadata,
            "notebook": self._extract_notebook_metadata,
            "image": self._extract_image_metadata,
            "scientific": self._extract_scientific_metadata,
            "data": self._extract_data_metadata,
        }

        extractor = extractors.get(category)
        if extractor:
            try:
                return extractor(file_path, ext)
            except Exception as e:
                return {"extraction_error": f"{category}: {str(e)[:50]}"}

        return {}

    def _extract_table_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from CSV/Excel files."""
        metadata: dict[str, Any] = {}

        if ext == ".csv" or ext == ".tsv":
            return self._extract_csv_metadata(file_path, ext)
        elif ext in (".xlsx", ".xls", ".xlsm", ".xltx"):
            return self._extract_excel_metadata(file_path)

        return metadata

    def _extract_csv_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from CSV/TSV files."""
        metadata: dict[str, Any] = {}

        try:
            import pandas as pd

            delimiter = "\t" if ext == ".tsv" else ","

            # Read with pandas, handle errors gracefully
            df = pd.read_csv(
                file_path,
                delimiter=delimiter,
                nrows=1000,  # Limit rows for performance
                on_bad_lines="skip",
                encoding_errors="ignore",
            )

            metadata["row_count"] = len(df)
            metadata["column_count"] = len(df.columns)
            metadata["column_names"] = df.columns.tolist()[:20]  # Limit to 20

            # Analyze column types
            column_types = {}
            numeric_cols = []
            date_cols = []
            categorical_cols = []

            for col in df.columns[:20]:
                dtype = str(df[col].dtype)
                if "int" in dtype or "float" in dtype:
                    column_types[col] = "numeric"
                    numeric_cols.append(col)
                elif "datetime" in dtype:
                    column_types[col] = "datetime"
                    date_cols.append(col)
                elif "object" in dtype:
                    # Check if it might be a date
                    sample = df[col].dropna().head(5).astype(str)
                    if any(self._looks_like_date(s) for s in sample):
                        column_types[col] = "datetime"
                        date_cols.append(col)
                    elif df[col].nunique() < len(df) * 0.5:
                        column_types[col] = "categorical"
                        categorical_cols.append(col)
                    else:
                        column_types[col] = "text"
                else:
                    column_types[col] = dtype

            metadata["column_types"] = column_types
            metadata["numeric_columns"] = numeric_cols[:10]
            metadata["date_columns"] = date_cols[:5]
            metadata["categorical_columns"] = categorical_cols[:10]

            # Missing value ratio
            total_cells = df.shape[0] * df.shape[1]
            if total_cells > 0:
                missing_ratio = df.isna().sum().sum() / total_cells
                metadata["missing_value_ratio"] = round(missing_ratio, 3)

        except ImportError:
            metadata["extraction_note"] = "pandas not available"
        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_excel_metadata(self, file_path: str) -> dict[str, Any]:
        """Extract metadata from Excel files."""
        metadata: dict[str, Any] = {}

        try:
            import openpyxl

            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

            metadata["sheet_count"] = len(wb.sheetnames)
            metadata["sheet_names"] = wb.sheetnames[:10]

            # Analyze first sheet
            if wb.sheetnames:
                ws = wb[wb.sheetnames[0]]
                metadata["column_count"] = ws.max_column
                metadata["row_count"] = ws.max_row

                # Get column names from first row
                col_names = []
                for col in range(1, min(ws.max_column + 1, 21)):
                    cell_value = ws.cell(row=1, column=col).value
                    if cell_value:
                        col_names.append(str(cell_value))
                metadata["column_names"] = col_names

            # Check for formulas (sample)
            has_formulas = False
            if wb.sheetnames:
                wb_with_formulas = openpyxl.load_workbook(
                    file_path, read_only=True, data_only=False
                )
                ws_formula = wb_with_formulas[wb_with_formulas.sheetnames[0]]
                for row in ws_formula.iter_rows(max_row=50, max_col=20):
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and cell.value.startswith("="):
                            has_formulas = True
                            break
                    if has_formulas:
                        break
                wb_with_formulas.close()

            metadata["has_formulas"] = has_formulas
            wb.close()

        except ImportError:
            metadata["extraction_note"] = "openpyxl not available"
        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_document_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from document files (PDF, DOCX, etc.)."""
        metadata: dict[str, Any] = {}

        if ext == ".pdf":
            return self._extract_pdf_metadata(file_path)
        elif ext == ".docx":
            return self._extract_docx_metadata(file_path)
        elif ext in (".txt", ".md"):
            return self._extract_text_metadata(file_path)

        return metadata

    def _extract_pdf_metadata(self, file_path: str) -> dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata: dict[str, Any] = {}

        # Try PyMuPDF (fitz) first
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            metadata["page_count"] = len(doc)

            # Check for images and tables in first few pages
            has_images = False
            has_tables = False
            text_sample = ""

            for page_num in range(min(5, len(doc))):
                page = doc[page_num]
                images = page.get_images()
                if images:
                    has_images = True

                # Simple table detection via text blocks
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            spans = line.get("spans", [])
                            if len(spans) > 3:  # Multiple columns might indicate table
                                has_tables = True

                if page_num == 0:
                    text_sample = page.get_text()[:500]

            metadata["has_images"] = has_images
            metadata["has_tables"] = has_tables

            # Check for form fields
            if doc.is_form_pdf:
                metadata["is_fillable"] = True
                # Count form fields
                field_count = 0
                for page in doc:
                    widgets = page.widgets()
                    if widgets:
                        field_count += len(list(widgets))
                metadata["field_count"] = field_count

            doc.close()
            return metadata

        except ImportError:
            pass

        # Fallback to pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                # Sample first pages for tables
                has_tables = False
                for page in pdf.pages[:5]:
                    if page.extract_tables():
                        has_tables = True
                        break
                metadata["has_tables"] = has_tables

        except ImportError:
            metadata["extraction_note"] = "No PDF library available (fitz/pdfplumber)"
        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_docx_metadata(self, file_path: str) -> dict[str, Any]:
        """Extract metadata from DOCX files."""
        metadata: dict[str, Any] = {}

        try:
            from docx import Document

            doc = Document(file_path)

            # Count paragraphs and estimate page count
            para_count = len(doc.paragraphs)
            metadata["paragraph_count"] = para_count
            metadata["page_count_estimate"] = max(1, para_count // 30)

            # Extract section titles (headings)
            section_titles = []
            for para in doc.paragraphs:
                if para.style and para.style.name.startswith("Heading"):
                    if para.text.strip():
                        section_titles.append(para.text.strip())

            metadata["section_count"] = len(section_titles)
            metadata["section_titles"] = section_titles[:10]

            # Check for tables and images
            metadata["has_tables"] = len(doc.tables) > 0
            metadata["table_count"] = len(doc.tables)

            # Check for images (inline shapes)
            has_images = False
            for para in doc.paragraphs:
                if para._element.xpath(".//a:blip"):
                    has_images = True
                    break
            metadata["has_images"] = has_images

        except ImportError:
            metadata["extraction_note"] = "python-docx not available"
        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_text_metadata(self, file_path: str) -> dict[str, Any]:
        """Extract metadata from plain text/markdown files."""
        metadata: dict[str, Any] = {}

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            lines = content.split("\n")
            metadata["line_count"] = len(lines)
            metadata["word_count_estimate"] = len(content.split())

            # For markdown, extract headings
            ext = Path(file_path).suffix.lower()
            if ext == ".md":
                headings = [
                    line.lstrip("#").strip()
                    for line in lines
                    if line.startswith("#") and line.lstrip("#").strip()
                ]
                metadata["section_titles"] = headings[:10]
                metadata["section_count"] = len(headings)

                # Check for code blocks
                metadata["has_code_blocks"] = "```" in content

        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_code_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from code files."""
        metadata: dict[str, Any] = {}

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            lines = content.split("\n")
            metadata["line_count"] = len(lines)

            # Language detection based on extension
            lang_map = {
                ".py": "python",
                ".js": "javascript",
                ".ts": "typescript",
                ".java": "java",
                ".cpp": "cpp",
                ".c": "c",
                ".go": "go",
                ".rs": "rust",
                ".rb": "ruby",
                ".php": "php",
                ".sh": "bash",
            }
            metadata["programming_language"] = lang_map.get(ext, "unknown")

            # Python-specific analysis
            if ext == ".py":
                try:
                    tree = ast.parse(content)

                    # Count functions and classes
                    functions = [
                        node.name
                        for node in ast.walk(tree)
                        if isinstance(node, ast.FunctionDef)
                    ]
                    classes = [
                        node.name
                        for node in ast.walk(tree)
                        if isinstance(node, ast.ClassDef)
                    ]

                    metadata["function_count"] = len(functions)
                    metadata["function_names"] = functions[:10]
                    metadata["class_count"] = len(classes)
                    metadata["class_names"] = classes[:5]

                    # Extract imports
                    imports = []
                    for node in ast.walk(tree):
                        if isinstance(node, ast.Import):
                            for alias in node.names:
                                imports.append(alias.name.split(".")[0])
                        elif isinstance(node, ast.ImportFrom):
                            if node.module:
                                imports.append(node.module.split(".")[0])

                    metadata["imports"] = list(set(imports))[:15]

                except SyntaxError:
                    # Simple regex fallback for imports
                    import re

                    imports = re.findall(r"^(?:from|import)\s+(\w+)", content, re.MULTILINE)
                    metadata["imports"] = list(set(imports))[:15]

        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_notebook_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from Jupyter notebooks."""
        metadata: dict[str, Any] = {}

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                nb = json.load(f)

            cells = nb.get("cells", [])
            metadata["cell_count"] = len(cells)

            # Count by cell type
            code_cells = [c for c in cells if c.get("cell_type") == "code"]
            markdown_cells = [c for c in cells if c.get("cell_type") == "markdown"]

            metadata["code_cell_count"] = len(code_cells)
            metadata["markdown_cell_count"] = len(markdown_cells)

            # Check for outputs
            has_outputs = any(c.get("outputs") for c in code_cells)
            metadata["has_outputs"] = has_outputs

            # Extract imports from code cells
            imports = set()
            for cell in code_cells:
                source = "".join(cell.get("source", []))
                import re

                found = re.findall(r"^(?:from|import)\s+(\w+)", source, re.MULTILINE)
                imports.update(found)

            metadata["imports"] = list(imports)[:15]

            # Get kernel info
            kernel = nb.get("metadata", {}).get("kernelspec", {})
            metadata["programming_language"] = kernel.get("language", "python")

        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_image_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from image files."""
        metadata: dict[str, Any] = {}

        try:
            from PIL import Image

            with Image.open(file_path) as img:
                metadata["width"] = img.width
                metadata["height"] = img.height
                metadata["color_mode"] = img.mode

                # Get DPI if available
                dpi = img.info.get("dpi")
                if dpi:
                    metadata["dpi"] = dpi[0] if isinstance(dpi, tuple) else dpi

                # Determine image type based on content
                if img.mode in ("1", "L"):
                    metadata["image_type"] = "grayscale"
                elif img.mode == "P":
                    metadata["image_type"] = "indexed"
                else:
                    metadata["image_type"] = "color"

        except ImportError:
            metadata["extraction_note"] = "PIL not available"
        except Exception as e:
            metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_scientific_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from scientific data files."""
        metadata: dict[str, Any] = {}

        # Map extension to scientific domain
        domain_map = {
            ".fasta": "bioinformatics",
            ".fa": "bioinformatics",
            ".fastq": "bioinformatics",
            ".fq": "bioinformatics",
            ".vcf": "bioinformatics",
            ".sam": "bioinformatics",
            ".bam": "bioinformatics",
            ".fits": "astronomy",
            ".geojson": "geospatial",
        }

        metadata["scientific_domain"] = domain_map.get(ext, "unknown")
        metadata["data_format"] = ext.lstrip(".")

        # Attempt to count records
        if ext in (".fasta", ".fa"):
            try:
                with open(file_path, "r") as f:
                    sample_count = sum(1 for line in f if line.startswith(">"))
                metadata["sample_count"] = sample_count
            except Exception:
                pass

        elif ext in (".fastq", ".fq"):
            try:
                with open(file_path, "r") as f:
                    lines = sum(1 for _ in f)
                metadata["sample_count"] = lines // 4
            except Exception:
                pass

        elif ext == ".geojson":
            try:
                with open(file_path, "r") as f:
                    data = json.load(f)
                features = data.get("features", [])
                metadata["feature_count"] = len(features)

                # Get coordinate system
                crs = data.get("crs", {}).get("properties", {}).get("name")
                if crs:
                    metadata["coordinate_system"] = crs
            except Exception:
                pass

        return metadata

    def _extract_data_metadata(self, file_path: str, ext: str) -> dict[str, Any]:
        """Extract metadata from JSON/XML/YAML files."""
        metadata: dict[str, Any] = {}

        if ext == ".json":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)

                if isinstance(data, list):
                    metadata["data_structure"] = "array"
                    metadata["item_count"] = len(data)
                elif isinstance(data, dict):
                    metadata["data_structure"] = "object"
                    metadata["top_level_keys"] = list(data.keys())[:10]

            except Exception as e:
                metadata["extraction_error"] = str(e)[:100]

        elif ext in (".yaml", ".yml"):
            try:
                import yaml

                with open(file_path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f)

                if isinstance(data, list):
                    metadata["data_structure"] = "array"
                    metadata["item_count"] = len(data)
                elif isinstance(data, dict):
                    metadata["data_structure"] = "object"
                    metadata["top_level_keys"] = list(data.keys())[:10]

            except ImportError:
                metadata["extraction_note"] = "PyYAML not available"
            except Exception as e:
                metadata["extraction_error"] = str(e)[:100]

        elif ext == ".xml":
            try:
                import xml.etree.ElementTree as ET

                tree = ET.parse(file_path)
                root = tree.getroot()

                metadata["root_tag"] = root.tag
                metadata["child_count"] = len(list(root))

            except Exception as e:
                metadata["extraction_error"] = str(e)[:100]

        return metadata

    def _extract_semantic(
        self,
        file_path: str,
        summary: str,
        content_type: str,
        programmatic_metadata: dict[str, Any],
    ) -> dict[str, Any]:
        """Extract semantic metadata using LLM."""
        if not self.llm:
            return {}

        path = Path(file_path)

        # Format programmatic metadata for context
        prog_meta_str = json.dumps(programmatic_metadata, indent=2, ensure_ascii=False)

        prompt = LLM_METADATA_EXTRACTION_PROMPT.format(
            filename=path.name,
            extension=path.suffix,
            file_size=programmatic_metadata.get("file_size_bytes", "unknown"),
            content_type=content_type,
            summary=summary[:2000],  # Limit summary length
            programmatic_metadata=prog_meta_str[:1500],  # Limit metadata length
        )

        response = self.llm.generate(
            system_prompt="You are a metadata extraction expert. Extract structured metadata from files. Return only valid JSON.",
            user_prompt=prompt,
            temperature=0.2,
        )

        # Parse JSON from response
        try:
            # Try to extract JSON from response
            response = response.strip()

            # Handle markdown code blocks
            if "```json" in response:
                start = response.find("```json") + 7
                end = response.find("```", start)
                response = response[start:end].strip()
            elif "```" in response:
                start = response.find("```") + 3
                end = response.find("```", start)
                response = response[start:end].strip()

            semantic_metadata = json.loads(response)

            # Validate expected fields
            valid_fields = {
                "language",
                "domain",
                "document_type",
                "complexity_level",
                "key_entities",
                "potential_operations",
            }
            return {k: v for k, v in semantic_metadata.items() if k in valid_fields}

        except json.JSONDecodeError as e:
            return {"semantic_extraction_error": f"JSON parse error: {str(e)[:50]}"}

    def _looks_like_date(self, value: str) -> bool:
        """Heuristic check if a string looks like a date."""
        import re

        date_patterns = [
            r"\d{4}-\d{2}-\d{2}",  # 2024-01-01
            r"\d{2}/\d{2}/\d{4}",  # 01/01/2024
            r"\d{2}-\d{2}-\d{4}",  # 01-01-2024
            r"\d{2}-\w{3}-\d{2,4}",  # 01-Jan-24
        ]
        return any(re.match(pattern, value) for pattern in date_patterns)
